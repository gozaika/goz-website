from pathlib import Path
import subprocess

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "C-1_Restaurant_Partner_Sales_Sheet"
ASSETS = OUT / "assets"
ICONS = ROOT / "icons"
PY_DEPS_NODE = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "node" / "bin" / "node.exe"


COLORS = {
    "saffron": "FF6B35",
    "forest": "1A5C38",
    "teal": "194B4A",
    "gold": "D4A017",
    "cream": "FFF8F0",
    "charcoal": "2D2D2D",
    "mist": "F5F5F5",
    "white": "FFFFFF",
    "line": "E7D7C6",
}


COPY = {
    "en": {
        "filename": "goZaika_C-1_A4_Restaurant_Partner_Sales_Sheet_EN.docx",
        "font": "Inter",
        "headline_font": "Poppins",
        "script_font": "Hind",
        "hero": "A new customer.\nWalking through your door.",
        "sub": "With goZaika, curated BAM Bag drops become a smart customer acquisition channel while your brand stays premium and in your control.",
        "jingle": "BAM! बड़ा ज़ायका, आएगा मज़ा",
        "section_title": "Four reasons goZaika works for serious restaurant brands",
        "cards": [
            ("Rs", "Recover margin. Not reputation.", "Turn planned kitchen prep into a revenue channel without weakening your menu value. Just smart release timing."),
            ("CRM", "New customers. Your data.", "Every BAM Bag claim introduces a customer who chose you. Customer details flow back to your CRM, not into a black box."),
            ("IN", "They walk in. You sell more.", "41% of pickup customers buy additional items at collection. Every window is a counter upsell moment."),
            ("BRAND", "Your brand stays your brand.", "Your restaurant name, food story, and hospitality stay front and centre. You decide what drops, when it drops, and how it is presented."),
        ],
        "stats": [
            ("12%", "Commission", "vs. 25-30% on aggregators"),
            ("0%", "First 30 Days", "Pilot without platform commission"),
            ("₹35,000+", "Base-Case Monthly Net", "at 7 BAM Bags per day"),
            ("Pickup only", "No Delivery Ops", "customers collect at your counter"),
        ],
        "cta": "Ready to drop your first BAM Bag?",
        "contact": "gozaika.in/for-restaurants  |  partners@gozaika.in",
        "qr": "QR TO\nPARTNER PAGE",
    },
    "hi": {
        "filename": "goZaika_C-1_A4_Restaurant_Partner_Sales_Sheet_HI.docx",
        "font": "Hind",
        "headline_font": "Hind",
        "script_font": "Hind",
        "hero": "नया customer.\nसीधा आपके restaurant के अंदर.",
        "sub": "goZaika आपके brand की पहचान संभालते हुए BAM Bag drops को smart customer acquisition channel बनाता है.",
        "jingle": "BAM! बड़ा ज़ायका, आएगा मज़ा",
        "section_title": "Serious restaurant brands के लिए goZaika के चार मजबूत कारण",
        "cards": [
            ("Rs", "Margin वापस. पहचान intact.", "Planned kitchen prep को नई revenue stream बनाइए. Menu value और brand dignity आपके control में रहती है."),
            ("CRM", "नए customers. Data आपका.", "हर BAM Bag claim एक नया customer लाता है. Aapke customers, aapka data - hamesha."),
            ("IN", "वो walk-in करते हैं. आप और sell करते हैं.", "Pickup customers में से 41% collection पर extra items खरीदते हैं. हर pickup window counter upsell moment है."),
            ("BRAND", "आपका खाना, आपकी पहचान.", "Restaurant name, food story और hospitality front and centre रहती है. Kya drop hoga, kab hoga - aap decide karte hain."),
        ],
        "stats": [
            ("12%", "Commission", "Aggregators: 25-30% तक"),
            ("0%", "पहले 30 दिन", "Commission के बिना pilot"),
            ("₹35,000+", "Monthly net, base case", "7 BAM Bags/day पर"),
            ("Pickup only", "Delivery ops नहीं", "Customers आपके counter पर आते हैं"),
        ],
        "cta": "अपना पहला BAM Bag drop करने के लिए ready?",
        "contact": "gozaika.in/for-restaurants  |  partners@gozaika.in",
        "qr": "QR\nPARTNER PAGE",
    },
    "te": {
        "filename": "goZaika_C-1_A4_Restaurant_Partner_Sales_Sheet_TE.docx",
        "font": "Noto Sans Telugu",
        "headline_font": "Noto Sans Telugu",
        "script_font": "Hind",
        "hero": "కొత్త customer.\nనేరుగా మీ restaurant లోకి.",
        "sub": "goZaika తో, curated BAM Bag drops మీ brand premium గా ఉండేలా చూసుకుంటూ smart customer acquisition channel గా పనిచేస్తాయి.",
        "jingle": "BAM! बड़ा ज़ायका, आएगा मज़ा",
        "section_title": "Serious restaurant brands కోసం goZaika పని చేసే 4 కారణాలు",
        "cards": [
            ("Rs", "Margin recover చేయండి. Reputation అలాగే ఉంచండి.", "Planned kitchen prep ను కొత్త revenue channel గా మార్చండి. Menu value మీ control లోనే ఉంటుంది."),
            ("CRM", "కొత్త customers. Data మీదే.", "ప్రతి BAM Bag claim మీను ఎంచుకున్న customer ని తీసుకొస్తుంది. ప్రతి customer ను మీరు own చేసుకోవచ్చు."),
            ("IN", "వాళ్లు walk-in అవుతారు. మీరు మరింత sell చేస్తారు.", "Pickup customers లో 41% collection సమయంలో extra items కొనుగోలు చేస్తారు. ప్రతి pickup window ఒక upsell moment."),
            ("BRAND", "మీ restaurant brand మీదే.", "మీ restaurant పేరు, food story, hospitality front and centre లో ఉంటాయి. ఏది drop చేయాలి, ఎప్పుడు చేయాలి - మీరు decide చేస్తారు."),
        ],
        "stats": [
            ("12%", "Commission", "Aggregators లో 25-30% వరకు"),
            ("0%", "మొదటి 30 రోజులు", "Commission లేకుండా pilot"),
            ("₹35,000+", "Monthly net, base case", "7 BAM Bags/day వద్ద"),
            ("Pickup only", "Delivery ops లేదు", "Customers మీ counter కి వస్తారు"),
        ],
        "cta": "మీ మొదటి BAM Bag drop కి ready?",
        "contact": "gozaika.in/for-restaurants  |  partners@gozaika.in",
        "qr": "QR\nPARTNER PAGE",
    },
}


def hex_to_rgb(hex_value):
    return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))


def set_run(run, font, size, color="2D2D2D", bold=False, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = hex_to_rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_borders(cell, color="FFFFFF", size="0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{side}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single" if size != "0" else "nil")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_width(table, width_twips):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def clear_para(p, after=0, before=0, align=None, line=1.05):
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(cell, text, font, size, color, bold=False, italic=False, after=0, before=0, align=None, line=1.05):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    clear_para(p, after=after, before=before, align=align, line=line)
    for idx, part in enumerate(text.split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run(run, font, size, color, bold=bold, italic=italic)
    return p


def make_qr_placeholder(path):
    from reportlab.graphics.barcode import qr

    qr_code = qr.QrCodeWidget("https://gozaika.in/for-restaurants").qr
    qr_code.make()
    count = qr_code.getModuleCount()
    border = 4
    block = 420 // (count + border * 2)
    size = block * (count + border * 2)
    img = Image.new("RGB", (size, size), "white")
    draw = ImageDraw.Draw(img)
    for y in range(count):
        for x in range(count):
            if qr_code.isDark(y, x):
                x0 = (x + border) * block
                y0 = (y + border) * block
                draw.rectangle([x0, y0, x0 + block - 1, y0 + block - 1], fill=(26, 92, 56))
    draw.rectangle([0, 0, size - 1, size - 1], outline=(26, 92, 56), width=6)
    img.save(path)


def pil_font(name, size, bold=False, italic=False):
    fonts = Path("C:/Windows/Fonts")
    if name == "nirmala":
        return ImageFont.truetype(str(fonts / "Nirmala.ttc"), size)
    if name == "display":
        return ImageFont.truetype(str(fonts / "georgiab.ttf"), size)
    if bold and italic:
        return ImageFont.truetype(str(fonts / "arialbi.ttf"), size)
    if bold:
        return ImageFont.truetype(str(fonts / "arialbd.ttf"), size)
    if italic:
        return ImageFont.truetype(str(fonts / "ariali.ttf"), size)
    return ImageFont.truetype(str(fonts / "arial.ttf"), size)


def draw_wrapped(draw, text, xy, max_width, font, fill, spacing=6, align="left"):
    x, y = xy
    lines = []
    for raw_line in text.split("\n"):
        words = raw_line.split(" ")
        line = ""
        for word in words:
            candidate = word if not line else f"{line} {word}"
            width = draw.textbbox((0, 0), candidate, font=font)[2]
            if width <= max_width or not line:
                line = candidate
            else:
                lines.append(line)
                line = word
        lines.append(line)
    ascent, descent = font.getmetrics()
    line_step = int((ascent + descent) * 1.08) + spacing
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_width = bbox[2] - bbox[0]
        tx = x
        if align == "center":
            tx = x + (max_width - line_width) / 2
        elif align == "right":
            tx = x + max_width - line_width
        draw.text((tx, y), line, font=font, fill=fill)
        y += line_step
    return y


def paste_asset(canvas, asset, box):
    img = Image.open(asset).convert("RGBA")
    img.thumbnail((box[2] - box[0], box[3] - box[1]), Image.LANCZOS)
    x = box[0] + ((box[2] - box[0]) - img.width) // 2
    y = box[1] + ((box[3] - box[1]) - img.height) // 2
    canvas.alpha_composite(img, (x, y))


def create_sheet_image(lang, data):
    width, height = 2480, 3508
    img = Image.new("RGBA", (width, height), "#" + COLORS["cream"])
    draw = ImageDraw.Draw(img)
    font_key = "nirmala" if lang in {"hi", "te"} else "arial"
    headline_key = "display" if lang == "en" else "nirmala"
    body = pil_font(font_key, 47)
    body_small = pil_font(font_key, 40 if lang != "te" else 35)
    body_tiny = pil_font(font_key, 34)
    card_title = pil_font(font_key, 56 if lang == "en" else 50 if lang == "hi" else 46, bold=True)
    hero_font = pil_font(headline_key, 120 if lang == "en" else 90 if lang == "hi" else 86, bold=True)
    sub_font = pil_font(font_key, 46)
    label_font = pil_font("arial", 34, bold=True)
    stat_font = pil_font("arial", 63, bold=True)
    stat_label = pil_font(font_key, 36, bold=True)
    cta_font = pil_font(font_key, 58, bold=True)

    safe = 142
    hero_h = 875
    bottom_y = 2750
    draw.rectangle([0, 0, width, hero_h], fill="#" + COLORS["teal"])
    draw.rectangle([0, hero_h, width, bottom_y], fill="#" + COLORS["cream"])
    draw.rectangle([0, bottom_y, width, height], fill="#" + COLORS["forest"])

    paste_asset(img, ASSETS / "gozaika-logo-white.png", (safe, 80, safe + 480, 225))
    paste_asset(img, ASSETS / "flame.png", (width - safe - 150, 70, width - safe, 225))
    y = 310
    y = draw_wrapped(draw, data["hero"], (safe, y), 1600, hero_font, "#" + COLORS["white"], spacing=18)
    y = draw_wrapped(draw, data["sub"], (safe, y + 28), 1860, sub_font, "#FFF4E8", spacing=10)
    draw_wrapped(draw, data["jingle"], (safe, y + 22), 1500, pil_font("nirmala", 48), "#" + COLORS["gold"], spacing=8)

    draw.text((safe, 965), "Restaurant Partner Sales Sheet", font=label_font, fill="#" + COLORS["gold"])
    draw_wrapped(draw, data["section_title"], (safe, 1026), 1900, pil_font(font_key, 65 if lang == "en" else 55 if lang == "hi" else 50, bold=True), "#" + COLORS["forest"], spacing=8)

    card_w = (width - 2 * safe - 55) // 2
    card_h = 535
    card_positions = [
        (safe, 1195),
        (safe + card_w + 55, 1195),
        (safe, 1195 + card_h + 55),
        (safe + card_w + 55, 1195 + card_h + 55),
    ]
    for (x, y), (icon, title, text) in zip(card_positions, data["cards"]):
        draw.rounded_rectangle([x, y, x + card_w, y + card_h], radius=28, fill="white", outline="#" + COLORS["line"], width=4)
        draw.rounded_rectangle([x + 42, y + 42, x + 180, y + 108], radius=26, fill="#FFF0E8")
        draw.text((x + 58, y + 58), icon, font=pil_font("arial", 30, bold=True), fill="#" + COLORS["saffron"])
        title_end = draw_wrapped(draw, title, (x + 42, y + 140), card_w - 84, card_title, "#" + COLORS["forest"], spacing=8)
        draw_wrapped(draw, text, (x + 42, max(y + 250, title_end + 18)), card_w - 84, body_small, "#" + COLORS["charcoal"], spacing=8)

    stat_y = bottom_y + 92
    stat_w = (width - 2 * safe) // 4
    for i, stat in enumerate(data["stats"]):
        x = safe + i * stat_w
        if i:
            draw.line([x, stat_y - 18, x, stat_y + 350], fill="#2F7A50", width=4)
        draw_wrapped(draw, stat[0], (x + 14, stat_y), stat_w - 28, stat_font if i != 3 else pil_font("arial", 52, bold=True), "white", spacing=4, align="center")
        draw_wrapped(draw, stat[1], (x + 14, stat_y + 110), stat_w - 28, stat_label, "#" + COLORS["gold"], spacing=4, align="center")
        draw_wrapped(draw, stat[2], (x + 14, stat_y + 175), stat_w - 28, body_tiny, "#E8F4E8", spacing=4, align="center")

    paste_asset(img, ASSETS / "flame.png", (safe, height - 365, safe + 105, height - 245))
    draw_wrapped(draw, data["cta"], (safe + 130, height - 405), 1420, cta_font, "white", spacing=8)
    draw_wrapped(draw, data["contact"], (safe + 130, height - 275), 1420, body, "#E8F4E8", spacing=6)
    paste_asset(img, ASSETS / "qr-placeholder.png", (width - safe - 260, height - 405, width - safe, height - 145))

    png_path = OUT / data["filename"].replace(".docx", ".png")
    pdf_path = OUT / data["filename"].replace(".docx", ".pdf")
    img.convert("RGB").save(png_path, dpi=(300, 300))
    img.convert("RGB").save(pdf_path, "PDF", resolution=300.0)


def convert_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    script = (
        "const sharp=require('sharp');"
        f"sharp({str(ICONS / 'gozaika-logo-white.svg')!r}).resize(1200).png().toFile({str(ASSETS / 'gozaika-logo-white.png')!r});"
        f"sharp({str(ICONS / 'gozaika-logo-color.svg')!r}).resize(1200).png().toFile({str(ASSETS / 'gozaika-logo-color.png')!r});"
        f"sharp({str(ICONS / 'flame.svg')!r}).resize(512).png().toFile({str(ASSETS / 'flame.png')!r});"
    )
    subprocess.run([str(PY_DEPS_NODE), "-e", script], check=True)
    make_qr_placeholder(ASSETS / "qr-placeholder.png")


def create_doc(lang, data):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.9)
    section.left_margin = Cm(0.9)
    section.right_margin = Cm(0.9)
    section.header_distance = Cm(0.3)
    section.footer_distance = Cm(0.3)

    styles = doc.styles
    styles["Normal"].font.name = data["font"]
    styles["Normal"].font.size = Pt(9)

    main = doc.add_table(rows=3, cols=1)
    set_table_width(main, 10886)
    for row in main.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_borders(row.cells[0])
        set_cell_margins(row.cells[0], top=170, start=220, bottom=170, end=220)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    main.rows[0].height = Cm(6.7)
    main.rows[1].height = Cm(13.4)
    main.rows[2].height = Cm(6.0)

    hero = main.rows[0].cells[0]
    set_cell_shading(hero, COLORS["teal"])
    inner = hero.add_table(rows=1, cols=2)
    set_table_width(inner, 10440)
    set_col_widths(inner, [Cm(14.8), Cm(3.6)])
    for cell in inner.rows[0].cells:
        set_borders(cell)
        set_cell_shading(cell, COLORS["teal"])
        set_cell_margins(cell, 0, 0, 0, 0)
    p_logo = inner.cell(0, 0).paragraphs[0]
    clear_para(p_logo, after=6)
    p_logo.add_run().add_picture(str(ASSETS / "gozaika-logo-white.png"), width=Cm(4.25))
    p_flame = inner.cell(0, 1).paragraphs[0]
    clear_para(p_flame, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_flame.add_run().add_picture(str(ASSETS / "flame.png"), width=Cm(1.2))
    add_text(hero, data["hero"], data["headline_font"], 25 if lang == "en" else 22, COLORS["white"], bold=True, after=6, line=0.94)
    add_text(hero, data["sub"], data["font"], 9.8 if lang != "te" else 9.2, "FFF4E8", after=4, line=1.12)
    add_text(hero, data["jingle"], data["script_font"], 10.5, COLORS["gold"], bold=True, after=0)

    middle = main.rows[1].cells[0]
    set_cell_shading(middle, COLORS["cream"])
    add_text(middle, "Restaurant Partner Sales Sheet", "Poppins", 7.5, COLORS["gold"], bold=True, after=3)
    add_text(middle, data["section_title"], data["headline_font"], 14.5 if lang == "en" else 12.7, COLORS["forest"], bold=True, after=6, line=1.0)

    cards_table = middle.add_table(rows=2, cols=2)
    set_table_width(cards_table, 10440)
    set_col_widths(cards_table, [Cm(9.1), Cm(9.1)])
    for row in cards_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(4.45)
    for i, (icon, title, body) in enumerate(data["cards"]):
        cell = cards_table.cell(i // 2, i % 2)
        set_borders(cell, color=COLORS["line"], size="8")
        set_cell_shading(cell, COLORS["white"])
        set_cell_margins(cell, top=155, start=185, bottom=155, end=185)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(cell, icon, "Poppins", 7.2, COLORS["saffron"], bold=True, after=3)
        add_text(cell, title, data["headline_font"], 11.2 if lang != "te" else 10.3, COLORS["forest"], bold=True, after=3, line=1.02)
        add_text(cell, body, data["font"], 8.2 if lang != "te" else 7.45, COLORS["charcoal"], after=0, line=1.08)

    bottom = main.rows[2].cells[0]
    set_cell_shading(bottom, COLORS["forest"])
    stats_table = bottom.add_table(rows=1, cols=4)
    set_table_width(stats_table, 10440)
    set_col_widths(stats_table, [Cm(4.55), Cm(4.55), Cm(4.55), Cm(4.55)])
    for idx, stat in enumerate(data["stats"]):
        cell = stats_table.cell(0, idx)
        set_borders(cell, color="2F7A50", size="6")
        set_cell_shading(cell, COLORS["forest"])
        set_cell_margins(cell, top=95, start=100, bottom=95, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(cell, stat[0], "Poppins", 13.4 if idx != 3 else 11.2, COLORS["white"], bold=True, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(cell, stat[1], data["font"], 7.8, COLORS["gold"], bold=True, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(cell, stat[2], data["font"], 6.8 if lang != "te" else 6.2, "E8F4E8", after=0, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)

    cta = bottom.add_table(rows=1, cols=3)
    set_table_width(cta, 10440)
    set_col_widths(cta, [Cm(1.25), Cm(13.35), Cm(3.6)])
    for cell in cta.rows[0].cells:
        set_borders(cell)
        set_cell_shading(cell, COLORS["forest"])
        set_cell_margins(cell, top=90, start=40, bottom=40, end=40)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cta.cell(0, 0).paragraphs[0]
    clear_para(p, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.add_run().add_picture(str(ASSETS / "flame.png"), width=Cm(0.72))
    add_text(cta.cell(0, 1), data["cta"], data["headline_font"], 12.5 if lang == "en" else 11.2, COLORS["white"], bold=True, after=2)
    add_text(cta.cell(0, 1), data["contact"], data["font"], 8.1, "E8F4E8", after=0)
    p_qr = cta.cell(0, 2).paragraphs[0]
    clear_para(p_qr, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_qr.add_run().add_picture(str(ASSETS / "qr-placeholder.png"), width=Cm(1.75))

    doc.save(OUT / data["filename"])


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    convert_assets()
    for lang, data in COPY.items():
        create_doc(lang, data)
        create_sheet_image(lang, data)


if __name__ == "__main__":
    main()
